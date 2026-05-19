#!/usr/bin/env python3
"""
Spielscheune – Rechnungsgenerator für Kinderbetreuung
Erstellt PDF-Rechnungen aus abrechnungsdaten.csv und vorlage.docx.
"""

import json
import logging
import os
import smtplib
import subprocess
import sys
from email.message import EmailMessage
from datetime import datetime
from pathlib import Path

import openpyxl
import pandas as pd
from docx import Document

# ─── Konfiguration ────────────────────────────────────────────────────────────
BASE_DIR       = Path(__file__).parent
VORLAGE        = BASE_DIR / "vorlage.docx"
EXCEL_DATEI    = BASE_DIR / "spielscheune.xlsm"   # Fallback: .xlsx
HISTORIE_DATEI = BASE_DIR / "abrechnungshistorie.csv"
AUSGABE_ORDNER = BASE_DIR / "rechnungen"
NUMMER_DATEI   = BASE_DIR / "rechnungsnummern.json"
DUPLIKAT_DATEI = BASE_DIR / "bereits_abgerechnet.json"
LIBREOFFICE    = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

# ─── E-Mail-Konfiguration ─────────────────────────────────────────────────────
SMTP_SERVER = "smtp.strato.de"
SMTP_PORT   = 587
SMTP_USER   = "sonnenstunden@vialkowitsch.de"
SMTP_PASS   = os.environ.get("SMTP_PASS", "")   # Passwort aus .env laden

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)


# ─── Duplikat-Schutz ─────────────────────────────────────────────────────────

def _lade_abgerechnet() -> dict:
    if DUPLIKAT_DATEI.exists():
        return json.loads(DUPLIKAT_DATEI.read_text(encoding="utf-8"))
    return {}

def _speichere_abgerechnet(daten: dict) -> None:
    DUPLIKAT_DATEI.write_text(json.dumps(daten, indent=2, ensure_ascii=False), encoding="utf-8")

def bereits_abgerechnet(kundennummer: str, kind_vorname: str, lzr: str) -> str | None:
    """Gibt die frühere Rechnungsnummer zurück, wenn dieser Eintrag schon existiert."""
    key = f"{kundennummer}_{kind_vorname}_{lzr}"
    return _lade_abgerechnet().get(key)

def markiere_als_abgerechnet(kundennummer: str, kind_vorname: str, lzr: str, nummer: str) -> None:
    daten = _lade_abgerechnet()
    daten[f"{kundennummer}_{kind_vorname}_{lzr}"] = nummer
    _speichere_abgerechnet(daten)


# ─── Abrechnungshistorie & Frequenzprüfung ────────────────────────────────────

MONATE_DE = {
    "Januar": 1, "Februar": 2, "März": 3, "April": 4,
    "Mai": 5, "Juni": 6, "Juli": 7, "August": 8,
    "September": 9, "Oktober": 10, "November": 11, "Dezember": 12,
}
_MONATE_INV = {v: k for k, v in MONATE_DE.items()}

FREQUENZ_MONATE = {
    "monatlich": 1,
    "vierteljährlich": 3,
    "halbjährlich": 6,
    "jährlich": 12,
}

def _lzr_zu_ym(lzr: str) -> tuple[int, int]:
    name, jahr = lzr.strip().split()
    return int(jahr), MONATE_DE[name]

def _ym_zu_lzr(jahr: int, monat: int) -> str:
    return f"{_MONATE_INV[monat]} {jahr}"

def _addiere_monate(jahr: int, monat: int, n: int) -> tuple[int, int]:
    monat += n
    jahr  += (monat - 1) // 12
    monat  = ((monat - 1) % 12) + 1
    return jahr, monat

def lade_historie() -> dict[tuple[str, str], str]:
    """Gibt {(kundennummer, kind_vorname): letzter_lzr} zurück."""
    if not HISTORIE_DATEI.exists():
        return {}
    df = pd.read_csv(HISTORIE_DATEI, sep=";", encoding="utf-8-sig")
    return {
        (str(r["Kundennummer"]), str(r["Kind_Vorname"])): str(r["LZR"])
        for _, r in df.iterrows()
    }

def schreibe_historie(kundennummer: str, kind_vorname: str, lzr: str, nummer: str, datum: datetime) -> None:
    neu = not HISTORIE_DATEI.exists()
    with HISTORIE_DATEI.open("a", encoding="utf-8-sig", newline="") as f:
        if neu:
            f.write("Kundennummer;Kind_Vorname;LZR;Rechnungsnummer;Datum\n")
        f.write(f"{kundennummer};{kind_vorname};{lzr};{nummer};{datum.strftime('%d.%m.%Y')}\n")

def naechster_lzr(letzter_lzr: str, frequenz: str) -> str:
    """'April 2025' + 'vierteljährlich' → 'Juli 2025'"""
    j, m = _lzr_zu_ym(letzter_lzr)
    j, m = _addiere_monate(j, m, FREQUENZ_MONATE.get(frequenz, 1))
    return _ym_zu_lzr(j, m)

def ist_faellig(lzr: str, letzter_lzr: str, frequenz: str) -> tuple[bool, str]:
    """True wenn lzr >= nächster fälliger LZR; gibt auch den nächsten LZR zurück."""
    naechster = naechster_lzr(letzter_lzr, frequenz)
    return _lzr_zu_ym(lzr) >= _lzr_zu_ym(naechster), naechster


# ─── Rechnungsnummer ──────────────────────────────────────────────────────────

def naechste_nummer(datum: datetime) -> str:
    """Liefert die nächste eindeutige Rechnungsnummer und persistiert den Zähler."""
    key = datum.strftime("%Y%m")
    daten: dict = {}
    if NUMMER_DATEI.exists():
        daten = json.loads(NUMMER_DATEI.read_text(encoding="utf-8"))
    n = daten.get(key, 0) + 1
    daten[key] = n
    NUMMER_DATEI.write_text(json.dumps(daten, indent=2), encoding="utf-8")
    return f"RE-{key}-{n:04d}"


# ─── Platzhalter-Ersetzung ────────────────────────────────────────────────────
# Word teilt Platzhalter ({Key}) oft auf mehrere Runs auf, z. B.:
#   run[0] = "{"   run[1] = "Strasse"   run[2] = "}"
# _verbinde_geteilte_platzhalter fasst solche Runs zusammen, bevor wir ersetzen.

def _verbinde_geteilte_platzhalter(paragraph) -> None:
    runs = paragraph.runs
    i = 0
    while i < len(runs):
        text = runs[i].text
        pos = text.find("{")
        # If '{' exists in this run but has no matching '}' after it: merge forward
        if pos != -1 and "}" not in text[pos:]:
            j = i + 1
            combined = text
            while j < len(runs):
                combined += runs[j].text
                if "}" in runs[j].text:
                    break
                j += 1
            runs[i].text = combined
            for k in range(i + 1, min(j + 1, len(runs))):
                runs[k].text = ""
        i += 1


def _ersetze_in_absatz(paragraph, ersetzungen: dict) -> None:
    _verbinde_geteilte_platzhalter(paragraph)
    for run in paragraph.runs:
        for key, val in ersetzungen.items():
            if key in run.text:
                run.text = run.text.replace(key, str(val))


def ersetze_platzhalter(doc: Document, ersetzungen: dict) -> None:
    for para in doc.paragraphs:
        _ersetze_in_absatz(para, ersetzungen)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _ersetze_in_absatz(para, ersetzungen)


# ─── Hilfsfunktionen ─────────────────────────────────────────────────────────

def eur(wert: float) -> str:
    return f"{wert:.2f}".replace(".", ",")


def zu_pdf(docx_pfad: Path) -> Path:
    if not Path(LIBREOFFICE).exists():
        log.error("LibreOffice nicht gefunden: %s", LIBREOFFICE)
        sys.exit(1)
    result = subprocess.run(
        [LIBREOFFICE, "--headless", "--convert-to", "pdf",
         "--outdir", str(docx_pfad.parent), str(docx_pfad)],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice: {result.stderr.strip()}")
    return docx_pfad.with_suffix(".pdf")


def lade_csv(pfad: Path) -> pd.DataFrame:
    for enc in ("utf-8-sig", "latin-1"):
        try:
            df = pd.read_csv(pfad, sep=";", encoding=enc)
            df.columns = df.columns.str.strip()
            return df
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"CSV konnte nicht gelesen werden: {pfad}")


def lade_excel() -> tuple[pd.DataFrame, pd.DataFrame]:
    pfad = EXCEL_DATEI if EXCEL_DATEI.exists() else EXCEL_DATEI.with_suffix(".xlsx")
    if not pfad.exists():
        raise FileNotFoundError(f"Excel-Datei nicht gefunden: {EXCEL_DATEI}")

    wb = openpyxl.load_workbook(pfad, data_only=True, read_only=True)

    def ws_zu_df(name: str, drop: list[str] | None = None) -> pd.DataFrame:
        ws = wb[name]
        rows = [[c.value for c in row] for row in ws.iter_rows()]
        rows = [r for r in rows if any(v is not None for v in r)]
        if len(rows) < 2:
            return pd.DataFrame()
        df = pd.DataFrame(rows[1:], columns=rows[0])
        df = df.loc[:, df.columns.notna()]        # Button-Spalten (None-Header) entfernen
        df = df.dropna(subset=[df.columns[0]])    # leere Zeilen am Ende entfernen
        if drop:
            df = df.drop(columns=drop, errors="ignore")
        return df

    kunden     = ws_zu_df("Kundenkartei")
    abrechnung = ws_zu_df("Monatsabrechnung", drop=["Name"])
    wb.close()
    return kunden, abrechnung


# ─── Rechnung erstellen ───────────────────────────────────────────────────────

def erstelle_rechnung(row: pd.Series, nummer: str, datum: datetime) -> Path:
    menge  = int(row["Anzahl_Essen"])
    preis  = float(str(row["Preis_pro_Essen"]).replace(",", "."))
    gesamt = menge * preis

    ersetzungen = {
        "{Name}":                row["Name"],
        "{Strasse}":             row["Strasse"],
        "{PLZ}":                 str(row["PLZ"]),
        "{Ort}":                 row["Ort"],
        "{Kundennummer}":        str(row["Kundennummer"]),
        "{Abrechnungsfrequenz}": row["Abrechnungsfrequenz"],
        "{Kind_Vorname}":        row["Kind_Vorname"],
        "{LZR}":                 row["LZR"],
        "{Rechnungsnummer}":     nummer,
        "{Rechnungsdatum}":      datum.strftime("%d.%m.%Y"),
        "{Anzahl_Essen}":        str(menge),
        "{Preis_pro_Essen}":     eur(preis),
        "{Gesamtpreis}":         eur(gesamt),
    }

    doc = Document(VORLAGE)
    ersetze_platzhalter(doc, ersetzungen)

    # Ablage in rechnungen/YYYY-MM/
    monat_ordner = AUSGABE_ORDNER / datum.strftime("%Y-%m")
    monat_ordner.mkdir(parents=True, exist_ok=True)

    docx_pfad = monat_ordner / f"Rechnung_{row['Kundennummer']}_{nummer}.docx"
    doc.save(docx_pfad)

    pdf_pfad = zu_pdf(docx_pfad)
    log.info("  ✓  %-20s %s  →  %s", row["Name"], nummer, pdf_pfad.name)
    return pdf_pfad


# ─── E-Mail-Versand ───────────────────────────────────────────────────────────

def sende_rechnung(empfaenger: str, pdf_pfad: Path, nummer: str, name: str) -> None:
    if not SMTP_PASS:
        raise RuntimeError("SMTP_PASS nicht gesetzt – bitte in .env eintragen")
    msg = EmailMessage()
    msg["Subject"] = f"Rechnung {nummer} – Spielscheune Sonnenstunden"
    msg["From"]    = SMTP_USER
    msg["To"]      = empfaenger
    msg.set_content(
        f"Liebe Eltern von {name},\n\n"
        "im Anhang erhalten Sie Ihre aktuelle Rechnung.\n\n"
        "Bei Fragen stehe ich gerne zur Verfügung.\n\n"
        "Herzliche Grüße\n"
        "Claudia Vialkowitsch\n"
        "Spielscheune Sonnenstunden"
    )
    with open(pdf_pfad, "rb") as f:
        msg.add_attachment(f.read(), maintype="application", subtype="pdf",
                           filename=pdf_pfad.name)
    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
        server.starttls()
        server.login(SMTP_USER, SMTP_PASS)
        server.send_message(msg)
    log.info("  ✉  Gesendet an %s", empfaenger)


# ─── Hauptprogramm ────────────────────────────────────────────────────────────

def main() -> None:
    log.info("Spielscheune Rechnungsgenerator")
    pfad = EXCEL_DATEI if EXCEL_DATEI.exists() else EXCEL_DATEI.with_suffix(".xlsx")
    log.info("Lade: %s", pfad.name)

    kunden, abrechnung = lade_excel()
    df = abrechnung.merge(kunden, on="Kundennummer", how="left", validate="m:1")
    fehlende = df["Name"].isna().sum()
    if fehlende:
        log.error("%d Zeile(n) ohne passenden Kundeneintrag – Abbruch", fehlende)
        sys.exit(1)
    datum = datetime.today()

    log.info("Datum: %s  |  %d Datensatz/Datensätze", datum.strftime("%d.%m.%Y"), len(df))
    fehler   = 0
    zukunft  = 0
    duplikat = 0
    erstellt = 0
    historie = lade_historie()

    for _, row in df.iterrows():
        kundennummer = str(row["Kundennummer"])
        lzr          = str(row["LZR"])
        kind_vorname = str(row["Kind_Vorname"])
        frequenz     = str(row["Abrechnungsfrequenz"])

        letzter = historie.get((kundennummer, kind_vorname))
        if letzter:
            faellig, naechster = ist_faellig(lzr, letzter, frequenz)
            if not faellig:
                log.info("  ⏭  %-20s noch nicht fällig – nächste: %s", kind_vorname, naechster)
                zukunft += 1
                continue

        vorherige = bereits_abgerechnet(kundennummer, kind_vorname, lzr)
        if vorherige:
            log.warning("  ⚠  %-20s bereits abgerechnet (%s) – übersprungen", kind_vorname, vorherige)
            duplikat += 1
            continue

        nummer = naechste_nummer(datum)
        try:
            pdf_pfad = erstelle_rechnung(row, nummer, datum)
            sende_rechnung(row["Rechnungsemail"], pdf_pfad, nummer, kind_vorname)
            markiere_als_abgerechnet(kundennummer, kind_vorname, lzr, nummer)
            schreibe_historie(kundennummer, kind_vorname, lzr, nummer, datum)
            historie[(kundennummer, kind_vorname)] = lzr
            erstellt += 1
        except Exception as exc:
            log.error("  ✗  %s: %s", row.get("Name", "?"), exc)
            fehler += 1

    teile = [f"{erstellt} erstellt"]
    if duplikat:
        teile.append(f"{duplikat} bereits abgerechnet")
    if zukunft:
        teile.append(f"{zukunft} noch nicht fällig")
    if fehler:
        teile.append(f"{fehler} Fehler")
    log.info("Fertig: %s  →  %s", ", ".join(teile), AUSGABE_ORDNER)
    sys.exit(1 if fehler else 0)


if __name__ == "__main__":
    main()

"""
Spielscheune – Rechnungsgenerator Tests
Deckt LZR-Parsing, Frequenzprüfung, Duplikatschutz, Rechnungsnummern,
Historie, Excel-Import und vollständige Integrationsläufe ab.
"""

import json
import sys
from datetime import datetime
from pathlib import Path

import openpyxl
import pytest

import CreateInvoice as CI


# ─── Hilfsfunktion: minimales Test-Excel anlegen ──────────────────────────────

def make_excel(base_path: Path, kunden: list, abrechnung: list) -> Path:
    pfad = base_path / "spielscheune.xlsx"
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    ws_k = wb.create_sheet("Kundenkartei")
    ws_k.append(["Kundennummer", "Name", "Strasse", "PLZ", "Ort",
                  "Rechnungsemail", "Abrechnungsfrequenz"])
    for row in kunden:
        ws_k.append(row)

    ws_a = wb.create_sheet("Monatsabrechnung")
    ws_a.append(["Kundennummer", "Name", "Kind_Vorname", "LZR",
                  "Anzahl_Essen", "Preis_pro_Essen"])
    for row in abrechnung:
        ws_a.append(row)

    wb.save(pfad)
    return pfad


# ─── Fixtures ─────────────────────────────────────────────────────────────────

@pytest.fixture(autouse=True)
def isoliere(tmp_path, monkeypatch):
    """Leitet jeden Datei-Zugriff in ein temporäres Verzeichnis um."""
    monkeypatch.setattr(CI, "BASE_DIR",       tmp_path)
    monkeypatch.setattr(CI, "DUPLIKAT_DATEI", tmp_path / "bereits_abgerechnet.json")
    monkeypatch.setattr(CI, "NUMMER_DATEI",   tmp_path / "rechnungsnummern.json")
    monkeypatch.setattr(CI, "HISTORIE_DATEI", tmp_path / "abrechnungshistorie.csv")
    monkeypatch.setattr(CI, "EXCEL_DATEI",    tmp_path / "spielscheune.xlsm")
    monkeypatch.setattr(CI, "AUSGABE_ORDNER", tmp_path / "rechnungen")
    monkeypatch.setattr(CI, "VORLAGE",        Path(__file__).parent / "vorlage.docx")
    return tmp_path


@pytest.fixture
def mock_pdf(monkeypatch):
    """Ersetzt LibreOffice durch eine Stub-Funktion, die eine Fake-PDF schreibt."""
    def fake_zu_pdf(docx_pfad):
        pdf = docx_pfad.with_suffix(".pdf")
        pdf.write_bytes(b"%PDF-1.4 fake")
        return pdf
    monkeypatch.setattr(CI, "zu_pdf", fake_zu_pdf)


@pytest.fixture
def mock_email(monkeypatch):
    """Unterdrückt SMTP-Versand vollständig."""
    monkeypatch.setattr(CI, "sende_rechnung", lambda *a, **kw: None)


# Standard-Testdaten (spiegeln das Testset aus der Entwicklung wider)
STD_KUNDEN = [
    ["K123", "Max Muster",     "Musterweg 1",     "12345", "Ort",           "test@test.de", "monatlich"],
    ["K124", "Lisa Beispiel",  "Beispielgasse 5", "54321", "Stadt",         "test@test.de", "vierteljährlich"],
    ["K125", "Tom Testeltern", "Teststraße 7",    "11223", "Beispielstadt", "test@test.de", "monatlich"],
]
STD_ABRECHNUNG = [
    ["K123", None, "Elias",  "April 2025", 20, 3.5],
    ["K124", None, "Jonas",  "April 2025", 60, 3.5],
    ["K125", None, "Jürgen", "April 2025", 18, 3.75],
    ["K125", None, "Emil",   "April 2025", 10, 3.5],
]


def run_main(tmp_path) -> int:
    """Führt main() aus und gibt den Exit-Code zurück."""
    with pytest.raises(SystemExit) as exc:
        CI.main()
    return exc.value.code


def setup_standard(tmp_path, monkeypatch, kunden=None, abrechnung=None):
    """Legt Standard-Excel an und setzt SMTP_PASS."""
    make_excel(tmp_path, kunden or STD_KUNDEN, abrechnung or STD_ABRECHNUNG)
    monkeypatch.setattr(CI, "SMTP_PASS", "fake-pass")


# ══════════════════════════════════════════════════════════════════════════════
# 1. LZR-Parsing und Monatsarithmetik
# ══════════════════════════════════════════════════════════════════════════════

class TestLzrParsing:
    @pytest.mark.parametrize("lzr,erwartet", [
        ("Januar 2025",   (2025,  1)),
        ("April 2025",    (2025,  4)),
        ("Dezember 2024", (2024, 12)),
        ("März 2026",     (2026,  3)),
    ])
    def test_lzr_zu_ym(self, lzr, erwartet):
        assert CI._lzr_zu_ym(lzr) == erwartet

    @pytest.mark.parametrize("monat", range(1, 13))
    def test_ym_zu_lzr_roundtrip(self, monat):
        lzr = CI._ym_zu_lzr(2025, monat)
        assert CI._lzr_zu_ym(lzr) == (2025, monat)

    @pytest.mark.parametrize("start_j,start_m,n,erwartet", [
        (2025,  1,  3, (2025,  4)),   # einfach vorwärts
        (2025, 11,  3, (2026,  2)),   # Jahreswechsel
        (2025, 12,  1, (2026,  1)),   # Dezember + 1
        (2025,  6, 12, (2026,  6)),   # genau 1 Jahr
        (2024,  1,  3, (2024,  4)),   # Schaltjahr-neutraler Bereich
    ])
    def test_addiere_monate(self, start_j, start_m, n, erwartet):
        assert CI._addiere_monate(start_j, start_m, n) == erwartet


# ══════════════════════════════════════════════════════════════════════════════
# 2. Frequenzprüfung
# ══════════════════════════════════════════════════════════════════════════════

class TestIstFaellig:
    def test_monatlich_naechster_monat_faellig(self):
        faellig, naechster = CI.ist_faellig("Mai 2025", "April 2025", "monatlich")
        assert faellig is True
        assert naechster == "Mai 2025"

    def test_monatlich_gleicher_monat_nicht_faellig(self):
        faellig, naechster = CI.ist_faellig("April 2025", "April 2025", "monatlich")
        assert faellig is False
        assert naechster == "Mai 2025"

    def test_monatlich_ueberfaellig_trotzdem_faellig(self):
        """Wenn man einen Monat vergisst, ist der übernächste Monat trotzdem fällig."""
        faellig, _ = CI.ist_faellig("Juni 2025", "April 2025", "monatlich")
        assert faellig is True

    def test_vierteljaehrlich_nach_einem_monat_nicht_faellig(self):
        faellig, naechster = CI.ist_faellig("Mai 2025", "April 2025", "vierteljährlich")
        assert faellig is False
        assert naechster == "Juli 2025"

    def test_vierteljaehrlich_nach_zwei_monaten_nicht_faellig(self):
        faellig, naechster = CI.ist_faellig("Juni 2025", "April 2025", "vierteljährlich")
        assert faellig is False
        assert naechster == "Juli 2025"

    def test_vierteljaehrlich_nach_drei_monaten_faellig(self):
        faellig, _ = CI.ist_faellig("Juli 2025", "April 2025", "vierteljährlich")
        assert faellig is True

    def test_vierteljaehrlich_ueber_jahreswechsel(self):
        faellig, naechster = CI.ist_faellig("Januar 2026", "Oktober 2025", "vierteljährlich")
        assert faellig is True
        faellig2, _ = CI.ist_faellig("Dezember 2025", "Oktober 2025", "vierteljährlich")
        assert faellig2 is False

    def test_halbjaehrlich_faellig(self):
        faellig, _ = CI.ist_faellig("Oktober 2025", "April 2025", "halbjährlich")
        assert faellig is True

    def test_halbjaehrlich_nicht_faellig(self):
        faellig, naechster = CI.ist_faellig("September 2025", "April 2025", "halbjährlich")
        assert faellig is False
        assert naechster == "Oktober 2025"

    def test_jaehrlich_faellig(self):
        faellig, _ = CI.ist_faellig("April 2026", "April 2025", "jährlich")
        assert faellig is True

    def test_jaehrlich_nicht_faellig(self):
        faellig, naechster = CI.ist_faellig("März 2026", "April 2025", "jährlich")
        assert faellig is False
        assert naechster == "April 2026"


# ══════════════════════════════════════════════════════════════════════════════
# 3. Duplikatschutz
# ══════════════════════════════════════════════════════════════════════════════

class TestDuplikatSchutz:
    def test_kein_eintrag_gibt_none(self):
        assert CI.bereits_abgerechnet("K123", "Elias", "April 2025") is None

    def test_markieren_und_wiederfinden(self):
        CI.markiere_als_abgerechnet("K123", "Elias", "April 2025", "RE-202505-0001")
        assert CI.bereits_abgerechnet("K123", "Elias", "April 2025") == "RE-202505-0001"

    def test_anderes_kind_gleiche_kundennummer_kein_duplikat(self):
        """Kernfall: zwei Kinder unter K125 dürfen sich nicht gegenseitig blockieren."""
        CI.markiere_als_abgerechnet("K125", "Jürgen", "April 2025", "RE-0001")
        assert CI.bereits_abgerechnet("K125", "Emil", "April 2025") is None

    def test_anderer_lzr_kein_duplikat(self):
        CI.markiere_als_abgerechnet("K123", "Elias", "April 2025", "RE-0001")
        assert CI.bereits_abgerechnet("K123", "Elias", "Mai 2025") is None

    def test_mehrere_eintraege_persistiert(self):
        CI.markiere_als_abgerechnet("K123", "Elias", "April 2025", "RE-0001")
        CI.markiere_als_abgerechnet("K124", "Jonas", "April 2025", "RE-0002")
        assert CI.bereits_abgerechnet("K123", "Elias", "April 2025") == "RE-0001"
        assert CI.bereits_abgerechnet("K124", "Jonas", "April 2025") == "RE-0002"


# ══════════════════════════════════════════════════════════════════════════════
# 4. Rechnungsnummern
# ══════════════════════════════════════════════════════════════════════════════

class TestRechnungsnummern:
    def test_erste_nummer(self):
        n = CI.naechste_nummer(datetime(2025, 5, 1))
        assert n == "RE-202505-0001"

    def test_fortlaufend_im_selben_monat(self):
        d = datetime(2025, 5, 1)
        assert CI.naechste_nummer(d) == "RE-202505-0001"
        assert CI.naechste_nummer(d) == "RE-202505-0002"
        assert CI.naechste_nummer(d) == "RE-202505-0003"

    def test_verschiedene_monate_unabhaengig(self):
        n_mai = CI.naechste_nummer(datetime(2025, 5, 1))
        n_jun = CI.naechste_nummer(datetime(2025, 6, 1))
        n_mai2 = CI.naechste_nummer(datetime(2025, 5, 1))
        assert n_mai  == "RE-202505-0001"
        assert n_jun  == "RE-202506-0001"
        assert n_mai2 == "RE-202505-0002"

    def test_format_vierstellig_gepaddet(self):
        d = datetime(2025, 5, 1)
        for _ in range(9):
            CI.naechste_nummer(d)
        n = CI.naechste_nummer(d)
        assert n == "RE-202505-0010"


# ══════════════════════════════════════════════════════════════════════════════
# 5. Abrechnungshistorie
# ══════════════════════════════════════════════════════════════════════════════

class TestHistorie:
    def test_leere_historie(self):
        assert CI.lade_historie() == {}

    def test_eintrag_schreiben_und_lesen(self):
        CI.schreibe_historie("K123", "Elias", "April 2025", "RE-0001", datetime(2025, 5, 1))
        h = CI.lade_historie()
        assert h[("K123", "Elias")] == "April 2025"

    def test_letzter_eintrag_gewinnt(self):
        CI.schreibe_historie("K123", "Elias", "April 2025", "RE-0001", datetime(2025, 5, 1))
        CI.schreibe_historie("K123", "Elias", "Mai 2025",   "RE-0005", datetime(2025, 6, 1))
        h = CI.lade_historie()
        assert h[("K123", "Elias")] == "Mai 2025"

    def test_zwei_kinder_gleiche_kundennummer(self):
        CI.schreibe_historie("K125", "Jürgen", "April 2025", "RE-0001", datetime(2025, 5, 1))
        CI.schreibe_historie("K125", "Emil",   "April 2025", "RE-0002", datetime(2025, 5, 1))
        h = CI.lade_historie()
        assert h[("K125", "Jürgen")] == "April 2025"
        assert h[("K125", "Emil")]   == "April 2025"

    def test_csv_hat_kopfzeile(self):
        CI.schreibe_historie("K123", "Elias", "April 2025", "RE-0001", datetime(2025, 5, 1))
        zeilen = CI.HISTORIE_DATEI.read_text(encoding="utf-8-sig").splitlines()
        assert zeilen[0] == "Kundennummer;Kind_Vorname;LZR;Rechnungsnummer;Datum"


# ══════════════════════════════════════════════════════════════════════════════
# 6. Excel-Import
# ══════════════════════════════════════════════════════════════════════════════

class TestLadeExcel:
    def test_kundenkartei_spalten_und_werte(self, tmp_path):
        make_excel(tmp_path, STD_KUNDEN, STD_ABRECHNUNG)
        k, _ = CI.lade_excel()
        assert list(k.columns) == [
            "Kundennummer", "Name", "Strasse", "PLZ",
            "Ort", "Rechnungsemail", "Abrechnungsfrequenz",
        ]
        assert len(k) == 3
        assert k.iloc[0]["Kundennummer"] == "K123"
        assert k.iloc[0]["Name"] == "Max Muster"

    def test_name_spalte_aus_abrechnung_entfernt(self, tmp_path):
        """Die XLOOKUP-Anzeigespalte darf nicht in den DataFrame gelangen."""
        make_excel(tmp_path, STD_KUNDEN, STD_ABRECHNUNG)
        _, a = CI.lade_excel()
        assert "Name" not in a.columns
        assert "Kundennummer" in a.columns

    def test_abrechnung_spalten(self, tmp_path):
        make_excel(tmp_path, STD_KUNDEN, STD_ABRECHNUNG)
        _, a = CI.lade_excel()
        assert list(a.columns) == [
            "Kundennummer", "Kind_Vorname", "LZR", "Anzahl_Essen", "Preis_pro_Essen",
        ]

    def test_leere_zeilen_ignoriert(self, tmp_path):
        pfad = tmp_path / "spielscheune.xlsx"
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        ws_k = wb.create_sheet("Kundenkartei")
        ws_k.append(["Kundennummer", "Name", "Strasse", "PLZ", "Ort", "Rechnungsemail", "Abrechnungsfrequenz"])
        ws_k.append(["K123", "Max Muster", "Weg", "12345", "Ort", "t@t.de", "monatlich"])
        ws_k.append([None] * 7)  # leere Zeile
        ws_a = wb.create_sheet("Monatsabrechnung")
        ws_a.append(["Kundennummer", "Name", "Kind_Vorname", "LZR", "Anzahl_Essen", "Preis_pro_Essen"])
        ws_a.append(["K123", None, "Elias", "April 2025", 20, 3.5])
        wb.save(pfad)
        k, _ = CI.lade_excel()
        assert len(k) == 1

    def test_preis_als_float_lesbar(self, tmp_path):
        make_excel(tmp_path, STD_KUNDEN, [["K123", None, "Elias", "April 2025", 20, 3.75]])
        _, a = CI.lade_excel()
        assert a.iloc[0]["Preis_pro_Essen"] == 3.75

    def test_fallback_auf_xlsx(self, tmp_path):
        """EXCEL_DATEI zeigt auf .xlsm (existiert nicht) → Fallback auf .xlsx."""
        make_excel(tmp_path, STD_KUNDEN, STD_ABRECHNUNG)
        # .xlsm existiert nicht → lade_excel() nutzt .xlsx
        k, a = CI.lade_excel()
        assert len(k) == 3
        assert len(a) == 4

    def test_fehlende_excel_wirft_fehler(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            CI.lade_excel()


# ══════════════════════════════════════════════════════════════════════════════
# 7. EUR-Formatierung
# ══════════════════════════════════════════════════════════════════════════════

class TestEur:
    @pytest.mark.parametrize("wert,erwartet", [
        (3.5,   "3,50"),
        (3.0,   "3,00"),
        (70.0,  "70,00"),
        (210.0, "210,00"),
        (0.0,   "0,00"),
    ])
    def test_eur_formatierung(self, wert, erwartet):
        assert CI.eur(wert) == erwartet


# ══════════════════════════════════════════════════════════════════════════════
# 8. Integrationstests – voller Abrechnungslauf
# ══════════════════════════════════════════════════════════════════════════════

class TestIntegration:

    def test_erster_lauf_alle_vier_rechnungen(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Erster Lauf: alle 4 Testeinträge werden abgerechnet."""
        setup_standard(tmp_path, monkeypatch)
        code = run_main(tmp_path)
        assert code == 0
        abgerechnet = json.loads(CI.DUPLIKAT_DATEI.read_text())
        assert len(abgerechnet) == 4

    def test_zwei_kinder_gleiche_kundennummer_unabhaengig(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """K125/Jürgen und K125/Emil erhalten separate Rechnungen mit verschiedenen Nummern."""
        setup_standard(tmp_path, monkeypatch)
        run_main(tmp_path)
        abgerechnet = json.loads(CI.DUPLIKAT_DATEI.read_text())
        re_juergen = abgerechnet["K125_Jürgen_April 2025"]
        re_emil    = abgerechnet["K125_Emil_April 2025"]
        assert re_juergen is not None
        assert re_emil    is not None
        assert re_juergen != re_emil

    def test_zweiter_lauf_erzeugt_keine_neuen_rechnungen(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Zweiter Lauf mit identischen Daten → Duplikatschutz greift, kein Zähler-Anstieg."""
        setup_standard(tmp_path, monkeypatch)
        run_main(tmp_path)
        zaehler_nach_lauf1 = json.loads(CI.NUMMER_DATEI.read_text())
        run_main(tmp_path)
        zaehler_nach_lauf2 = json.loads(CI.NUMMER_DATEI.read_text())
        assert zaehler_nach_lauf1 == zaehler_nach_lauf2

    def test_pdf_dateien_werden_angelegt(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Nach dem Lauf liegen genau 4 PDFs im Ausgabe-Ordner."""
        setup_standard(tmp_path, monkeypatch)
        run_main(tmp_path)
        pdfs = list(CI.AUSGABE_ORDNER.rglob("*.pdf"))
        assert len(pdfs) == 4

    def test_rechnungsnummern_fortlaufend_ohne_luecke(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Die 4 Rechnungsnummern enden auf 0001–0004 ohne Lücke."""
        setup_standard(tmp_path, monkeypatch)
        run_main(tmp_path)
        abgerechnet = json.loads(CI.DUPLIKAT_DATEI.read_text())
        suffixe = sorted(v.split("-")[-1] for v in abgerechnet.values())
        assert suffixe == ["0001", "0002", "0003", "0004"]

    def test_historie_nach_lauf_vollstaendig(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Alle 4 Kinder-LZR-Kombinationen sind in der Historie eingetragen."""
        setup_standard(tmp_path, monkeypatch)
        run_main(tmp_path)
        h = CI.lade_historie()
        assert h[("K123", "Elias")]  == "April 2025"
        assert h[("K124", "Jonas")]  == "April 2025"
        assert h[("K125", "Jürgen")] == "April 2025"
        assert h[("K125", "Emil")]   == "April 2025"

    def test_frequenz_vierteljaehrlich_zu_frueh_uebersprungen(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Lisa Beispiel (vierteljährlich): Mai-Abrechnung nach April-Lauf wird übersprungen."""
        kunden     = [["K124", "Lisa Beispiel", "Gasse 5", "54321", "Stadt", "t@t.de", "vierteljährlich"]]
        abr_april  = [["K124", None, "Jonas", "April 2025", 60, 3.5]]
        abr_mai    = [["K124", None, "Jonas", "Mai 2025",   60, 3.5]]

        make_excel(tmp_path, kunden, abr_april)
        monkeypatch.setattr(CI, "SMTP_PASS", "x")
        run_main(tmp_path)
        zaehler_nach_april = json.loads(CI.NUMMER_DATEI.read_text())

        make_excel(tmp_path, kunden, abr_mai)
        run_main(tmp_path)
        zaehler_nach_mai = json.loads(CI.NUMMER_DATEI.read_text())

        assert zaehler_nach_april == zaehler_nach_mai  # kein neuer Zähler-Anstieg

    def test_frequenz_vierteljaehrlich_nach_drei_monaten_abgerechnet(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Lisa Beispiel (vierteljährlich): Juli-Abrechnung nach April-Lauf wird durchgeführt."""
        kunden    = [["K124", "Lisa Beispiel", "Gasse 5", "54321", "Stadt", "t@t.de", "vierteljährlich"]]
        abr_april = [["K124", None, "Jonas", "April 2025", 60, 3.5]]
        abr_juli  = [["K124", None, "Jonas", "Juli 2025",  62, 3.5]]

        make_excel(tmp_path, kunden, abr_april)
        monkeypatch.setattr(CI, "SMTP_PASS", "x")
        run_main(tmp_path)
        zaehler_nach_april = json.loads(CI.NUMMER_DATEI.read_text())

        make_excel(tmp_path, kunden, abr_juli)
        run_main(tmp_path)
        zaehler_nach_juli = json.loads(CI.NUMMER_DATEI.read_text())

        monat_key = list(zaehler_nach_juli.keys())[0]
        assert zaehler_nach_juli[monat_key] > zaehler_nach_april[monat_key]

    def test_fehlender_kundeneintrag_exitcode_1(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Eintrag in Monatsabrechnung ohne passenden Kunden → Exit 1."""
        make_excel(tmp_path,
                   kunden=[["K123", "Max Muster", "Weg 1", "12345", "Ort", "t@t.de", "monatlich"]],
                   abrechnung=[["K999", None, "Elias", "April 2025", 20, 3.5]])
        monkeypatch.setattr(CI, "SMTP_PASS", "x")
        code = run_main(tmp_path)
        assert code == 1

    def test_gesamtpreis_in_docx_korrekt(self, tmp_path, monkeypatch, mock_email):
        """Prüft, dass der Gesamtpreis im DOCX korrekt berechnet und eingesetzt wird."""
        # 20 Essen × 3,50 € = 70,00 €
        kunden = [["K123", "Max Muster", "Weg 1", "12345", "Ort", "t@t.de", "monatlich"]]
        abr    = [["K123", None, "Elias", "April 2025", 20, 3.5]]
        make_excel(tmp_path, kunden, abr)
        monkeypatch.setattr(CI, "SMTP_PASS", "x")

        docx_gefunden = []
        original_zu_pdf = CI.zu_pdf

        def intercept_zu_pdf(docx_pfad):
            docx_gefunden.append(docx_pfad)
            pdf = docx_pfad.with_suffix(".pdf")
            pdf.write_bytes(b"%PDF fake")
            return pdf

        monkeypatch.setattr(CI, "zu_pdf", intercept_zu_pdf)
        run_main(tmp_path)

        from docx import Document
        assert docx_gefunden, "Kein DOCX wurde erzeugt"
        doc = Document(docx_gefunden[0])
        volltext = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    volltext += "\n" + "\n".join(p.text for p in cell.paragraphs)
        assert "70,00" in volltext

    def test_monatlich_nach_zwei_monaten_beide_abgerechnet(self, tmp_path, monkeypatch, mock_pdf, mock_email):
        """Monatlicher Kunde läuft zwei Monate hintereinander durch."""
        kunden    = [["K123", "Max Muster", "Weg 1", "12345", "Ort", "t@t.de", "monatlich"]]
        abr_april = [["K123", None, "Elias", "April 2025", 20, 3.5]]
        abr_mai   = [["K123", None, "Elias", "Mai 2025",   18, 3.5]]

        make_excel(tmp_path, kunden, abr_april)
        monkeypatch.setattr(CI, "SMTP_PASS", "x")
        run_main(tmp_path)

        make_excel(tmp_path, kunden, abr_mai)
        run_main(tmp_path)

        h = CI.lade_historie()
        assert h[("K123", "Elias")] == "Mai 2025"
        pdfs = list(CI.AUSGABE_ORDNER.rglob("*.pdf"))
        assert len(pdfs) == 2
